"""Generate SkillsHub_Phase1_Guide.docx.

Single-shot generator. Runs from the project root or backend dir; resolves the
output path to <repo_root>/docs/SkillsHub_Phase1_Guide.docx.
"""
from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

# ── Paths ──────────────────────────────────────────────────────────────
BACKEND_DIR = Path(__file__).resolve().parents[1]
REPO_ROOT = BACKEND_DIR.parent
OUT_DIR = REPO_ROOT / "docs"
OUT_DIR.mkdir(parents=True, exist_ok=True)
OUT_PATH = OUT_DIR / "SkillsHub_Phase1_Guide.docx"

# ── Style helpers ──────────────────────────────────────────────────────
BODY_FONT = "Calibri"
CODE_FONT = "Consolas"
BODY_SIZE = Pt(11)
H1_SIZE = Pt(16)
H2_SIZE = Pt(13)
H3_SIZE = Pt(12)
CODE_SIZE = Pt(9)
GRAY_BG = "F2F2F2"
HEADER_BG = "D9E1F2"


def _set_cell_shade(cell, color_hex: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def _set_cell_borders(cell) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "BFBFBF")
        tcBorders.append(b)
    tcPr.append(tcBorders)


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_heading("", level=level)
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = BODY_FONT
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x3B, 0x6E)
    if level == 1:
        run.font.size = H1_SIZE
    elif level == 2:
        run.font.size = H2_SIZE
    else:
        run.font.size = H3_SIZE


def add_para(doc: Document, text: str, *, bold: bool = False, italic: bool = False,
             align=None, size=None) -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = BODY_FONT
    run.font.size = size or BODY_SIZE
    run.bold = bold
    run.italic = italic


def add_bullet(doc: Document, text: str, *, level: int = 0) -> None:
    p = doc.add_paragraph(style="List Bullet")
    if level:
        p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = BODY_FONT
    run.font.size = BODY_SIZE


def add_numbered(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = BODY_FONT
    run.font.size = BODY_SIZE


def add_code_block(doc: Document, code: str) -> None:
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = True
    cell = tbl.rows[0].cells[0]
    _set_cell_shade(cell, GRAY_BG)
    _set_cell_borders(cell)
    cell.text = ""
    # Strip trailing whitespace on each line; keep blank lines.
    lines = [ln.rstrip() for ln in code.splitlines()]
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(line if line else " ")
        run.font.name = CODE_FONT
        run.font.size = CODE_SIZE
    # Tiny breathing room after the block.
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc: Document, headers: list[str], rows: list[list[str]],
              col_widths: list[float] | None = None) -> None:
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    hdr_cells = tbl.rows[0].cells
    for i, h in enumerate(headers):
        cell = hdr_cells[i]
        _set_cell_shade(cell, HEADER_BG)
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.name = BODY_FONT
        run.font.size = Pt(10)
    for row in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = BODY_FONT
            run.font.size = Pt(10)
    if col_widths:
        for row in tbl.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)


def add_page_break(doc: Document) -> None:
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    for tag, attrs, text in (
        ("w:fldChar", {"w:fldCharType": "begin"}, None),
        ("w:instrText", {"xml:space": "preserve"}, "PAGE"),
        ("w:fldChar", {"w:fldCharType": "end"}, None),
    ):
        el = OxmlElement(tag)
        for k, v in attrs.items():
            el.set(qn(k), v)
        if text is not None:
            el.text = text
        run._r.append(el)


def add_toc_field(paragraph) -> None:
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = (
        "Right-click here in Word and choose 'Update Field' to populate the table of contents."
    )
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for el in (fldChar1, instr, sep, placeholder, end):
        run._r.append(el)


# ── Document construction ─────────────────────────────────────────────
def build() -> None:
    doc = Document()

    # Margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Default body style
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = BODY_SIZE

    # Header (document title, right-aligned, small)
    header = doc.sections[0].header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = hp.add_run("SkillsHub — Phase 1 Project Guide")
    hr.font.name = BODY_FONT
    hr.font.size = Pt(8)
    hr.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    # Footer (page number, right-aligned)
    footer = doc.sections[0].footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_page_number(fp)
    # Style the page-number run
    for r in fp.runs:
        r.font.name = BODY_FONT
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # ── Cover page ────────────────────────────────────────────────
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("SkillsHub")
    r.font.name = BODY_FONT
    r.font.size = Pt(36)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x3B, 0x6E)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Phase 1 Project Guide")
    r.font.name = BODY_FONT
    r.font.size = Pt(22)
    r.font.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("AI-Powered Skills Intelligence Platform")
    r.font.name = BODY_FONT
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('"Type a question in plain English. Get a ranked, reasoned answer in seconds."')
    r.font.name = BODY_FONT
    r.font.size = Pt(12)
    r.italic = True

    for _ in range(6):
        doc.add_paragraph()

    meta = [
        ("Version", "1.0"),
        ("Date", date.today().isoformat()),
        ("Author", "KetanPatil27"),
        ("Repository", "VAST-MERN-Stack-Training / HACKATHON-PROJECT / skillshub-ai"),
    ]
    for label, value in meta:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"{label}: ")
        r.bold = True
        r.font.name = BODY_FONT
        r.font.size = Pt(11)
        r2 = p.add_run(value)
        r2.font.name = BODY_FONT
        r2.font.size = Pt(11)

    add_page_break(doc)

    # ── Table of Contents ─────────────────────────────────────────
    add_heading(doc, "Table of Contents", level=1)
    p = doc.add_paragraph()
    add_toc_field(p)
    add_para(
        doc,
        "Note: when this document opens for the first time, Word may show a "
        "placeholder above. Right-click the placeholder and choose Update Field "
        "(or press F9) to render the live, page-numbered TOC.",
        italic=True,
    )
    add_page_break(doc)

    # ─────────────────────────────────────────────────────────────
    # 1. EXECUTIVE SUMMARY
    # ─────────────────────────────────────────────────────────────
    add_heading(doc, "1. Executive Summary", level=1)

    add_para(
        doc,
        "SkillsHub is an AI-powered skills intelligence platform built for software "
        "companies that need to find the right person — fast. Today, when HR or a "
        "delivery manager needs someone who knows React and has shipped a payment "
        "integration, the answer lives in spreadsheets, in Slack DMs, and in the "
        "heads of ten different managers. SkillsHub replaces that scramble with one "
        "natural-language question and a ranked, evidence-backed shortlist.",
    )

    add_para(
        doc,
        "The product has two centerpieces. First, a resume upload that uses Google "
        "Gemini to turn a PDF into a structured profile — skills (with proficiency "
        "and evidence), projects, and AI-inferred related skills — in seconds. "
        "Second, a search box where HR types a real question (\"Senior frontend folks "
        "in Pune who haven't been on a new project this quarter\") and watches "
        "ranked candidate cards stream in, each with a match score, a plain-English "
        "reason, and the exact resume snippets that justify the match.",
    )

    add_para(
        doc,
        "Phase 1 ships a complete, demonstrable slice: JWT-secured auth with HR and "
        "Employee roles, AI extraction with inferred skills, a review queue, a "
        "filterable directory, three flavours of natural-language search (free-form, "
        "from a pasted JD, and a team-builder), and an audit log of every search "
        "for later analysis. Both demo accounts and 15 deterministically seeded "
        "profiles are wired up so the demo never depends on a live extraction.",
    )

    add_heading(doc, "Key Numbers", level=2)
    add_table(
        doc,
        ["Metric", "Value"],
        [
            ["Total backend Python files", "~45 (3,364 LOC)"],
            ["Total frontend TS/TSX files", "~50 (4,467 LOC)"],
            ["API endpoints", "21"],
            ["Database tables", "6"],
            ["AI prompt files (versioned)", "4 (extraction, inference, search, jd_distill)"],
            ["Seeded employee profiles", "15 (all APPROVED on seed)"],
            ["Average resume extraction time", "~4–8 s (Gemini 2.5 model dependent)"],
            ["Average search response (first card)", "~2–4 s (streamed; full ranking ~6–10 s)"],
            ["Showcase model", "gemini-2.5-pro (configurable)"],
            ["Light model (inference, JD distill)", "gemini-2.5-flash"],
        ],
        col_widths=[2.5, 3.5],
    )

    add_page_break(doc)

    # ─────────────────────────────────────────────────────────────
    # 2. PROBLEM & PURPOSE
    # ─────────────────────────────────────────────────────────────
    add_heading(doc, "2. Problem & Purpose", level=1)

    add_heading(doc, "2.1 The Problem", level=2)
    add_para(
        doc,
        "In a 200-person engineering organisation, basic staffing questions are "
        "surprisingly hard. \"Who in Pune knows Java and has shipped a Razorpay "
        "integration?\" sounds like a one-minute lookup, but the data needed to "
        "answer it lives in three places: stale skill spreadsheets, project "
        "post-mortem decks no one reads, and the institutional memory of senior "
        "managers who have to be pinged on Slack.",
    )
    add_para(
        doc,
        "The cost is real. A staffing decision that should take ten minutes can "
        "stretch into a week of follow-ups. Candidates already on the bench get "
        "missed because nobody remembered they did fintech work three projects "
        "ago. New joiners spend their first sprint introducing themselves to the "
        "same five team leads. Worse, the data that does exist is keyword-shaped — "
        "a search for \"payment gateway\" misses the engineer whose resume says "
        "\"integrated Razorpay checkout\" because the literal phrase is absent.",
    )
    add_para(
        doc,
        "The hackathon brief framed this as two problems: turn unstructured "
        "resumes into structured, queryable profiles, and let HR query those "
        "profiles in plain English. SkillsHub treats both as one design problem "
        "with a shared AI spine.",
    )

    add_heading(doc, "2.2 Our Purpose", level=2)
    add_para(
        doc,
        "SkillsHub aims to be the single source of truth for who-knows-what inside "
        "a software company — a place where every resume lands, gets parsed by AI, "
        "is checked by a human reviewer, and becomes immediately searchable in "
        "natural language. The product earns its place by being faster than the "
        "spreadsheets it replaces and more honest than the keyword search it "
        "out-thinks.",
    )
    add_para(doc, "Phase 1 success looks like:", bold=True)
    add_bullet(doc, "HR finds a qualified candidate from a plain-English question in under 30 seconds.")
    add_bullet(doc, "A resume PDF becomes an editable profile in under 10 seconds, with at least one inferred related skill.")
    add_bullet(doc, "Every search result cites at least one specific piece of evidence (project title, skill phrase, or domain) from the candidate's resume.")
    add_bullet(doc, "Every profile that lands in search has been seen and approved by a human — no AI output ships to a hiring manager unreviewed.")
    add_bullet(doc, "A new judge or teammate can run the full system locally in ten minutes flat, with two seeded one-click logins.")

    add_page_break(doc)

    # ─────────────────────────────────────────────────────────────
    # 3. HOW SKILLSHUB FULFILLS ITS PURPOSE
    # ─────────────────────────────────────────────────────────────
    add_heading(doc, "3. How SkillsHub Fulfills Its Purpose", level=1)

    add_para(
        doc,
        "Each success criterion above maps to a specific feature and a specific "
        "file in the repository. The table below is the index; the paragraphs that "
        "follow explain the design choice behind each row.",
    )

    add_table(
        doc,
        ["Goal", "How SkillsHub Delivers", "Where in Code"],
        [
            [
                "Candidate lookup in <30 s from plain English",
                "POST /search streams ranked cards via SSE; each card has a match score, reason, and evidence quotes",
                "backend/app/modules/search/router.py, service.py",
            ],
            [
                "Resume → editable profile in <10 s",
                "POST /resumes/upload runs PDF → text → Gemini extraction → related-skill inference → DB write",
                "backend/app/modules/resumes/service.py",
            ],
            [
                "Every result cites specific evidence",
                "Search prompt requires matched_skill_names and 1–3 evidence_snippets per candidate",
                "backend/app/modules/ai/prompts/search_v1.py",
            ],
            [
                "Human-in-the-loop before search visibility",
                "Profiles land in PENDING_REVIEW; HR reviews via /review/[id]; only APPROVED profiles enter the search corpus",
                "backend/app/modules/review/, search/service.py:list_approved_candidates",
            ],
            [
                "Ten-minute local boot for a new joiner",
                "alembic upgrade head + scripts/seed_database.py creates 15 approved profiles and 2 demo logins",
                "backend/scripts/seed_database.py, seed_data/profiles.json",
            ],
        ],
        col_widths=[1.7, 2.8, 1.9],
    )

    add_heading(doc, "Why each row is the right design", level=2)

    add_para(
        doc,
        "Streamed SSE results beat a blocking REST response for two reasons. First, "
        "the top result arrives in ~2 seconds, which feels live during a demo. "
        "Second, the model can output JSON elements one at a time, and the backend "
        "parses each complete object as soon as it arrives — no full-buffer wait. "
        "The endpoint lives in modules/search/router.py and the streaming logic in "
        "modules/ai/service.py:stream_search_results.",
    )
    add_para(
        doc,
        "The extraction pipeline is sequential by design: PDF parsing, then "
        "extraction, then inference. Inference is a separate AI call because it "
        "needs the extracted skills as input, and because it can fail "
        "non-fatally — if Gemini errors during inference, the resume still saves "
        "with explicit skills only. See resumes/service.py:upload_and_extract.",
    )
    add_para(
        doc,
        "Evidence-by-default in the search prompt is a deliberate guardrail "
        "against hallucination. The prompt forbids vague phrases like \"good fit\" "
        "and requires the model to copy exact strings from the candidate JSON. "
        "When a judge asks \"how do we know it's not making this up?\", we can "
        "point at the evidence_snippets field on every result card.",
    )
    add_para(
        doc,
        "The review queue keeps a human in the loop. A profile starts in "
        "PENDING_REVIEW, is rendered for HR with all extracted fields editable, "
        "and only becomes searchable when HR clicks Approve. The audit trail "
        "(reviewer_id, reviewer_notes, reviewed_at) is preserved on every "
        "ReviewQueueItem.",
    )
    add_para(
        doc,
        "The seeding script is idempotent: it upserts the two demo accounts, "
        "wipes any existing employee data, and re-creates 15 deterministic "
        "profiles. This means a demo accident — accidentally rejecting a profile, "
        "approving a junk record — is one command away from being undone.",
    )

    add_page_break(doc)

    # ─────────────────────────────────────────────────────────────
    # 4. TECH STACK
    # ─────────────────────────────────────────────────────────────
    add_heading(doc, "4. Tech Stack", level=1)

    add_para(
        doc,
        "Every choice in the stack was made against one constraint: build a "
        "complete, demoable system in two days while leaving a clean path to "
        "production. The table below shows what was picked and why, with exact "
        "versions read from requirements.txt and package.json.",
    )

    add_table(
        doc,
        ["Layer", "Technology", "Version", "Why we chose it"],
        [
            ["Frontend framework", "Next.js (App Router)", "14.2.16",
             "Native streaming support, route groups, edge middleware for auth. Faster than wiring SSE in CRA."],
            ["UI library", "shadcn/ui + Tailwind CSS", "tailwindcss 3.4.14",
             "Copy-paste primitives we can restyle without fighting a component framework. Zero runtime cost beyond Radix."],
            ["Animation", "Framer Motion", "11.11.10",
             "Result-card slide-ins and the multi-step extraction progress are what make the magic screens feel magical."],
            ["State / data", "TanStack Query", "5.59.16",
             "Cache invalidation handled for us. Used for everything that isn't an SSE stream."],
            ["Forms", "react-hook-form + Zod", "7.53.1 / 3.23.8",
             "Schema-validated forms with minimal re-renders. Same Zod schemas validate signup on the client and on the server through Pydantic."],
            ["Backend framework", "FastAPI", "0.115.6",
             "Auto-generated OpenAPI/Swagger lets judges click through every endpoint at /docs without a Postman setup."],
            ["ORM", "SQLAlchemy 2.0 async", "2.0.36",
             "Async-first, mature, plays well with asyncpg and Alembic."],
            ["Migrations", "Alembic", "1.14.0",
             "Standard tool for SQLAlchemy. One migration (0001_initial) covers the entire Phase 1 schema."],
            ["Database", "Supabase Postgres (asyncpg)", "asyncpg 0.30.0",
             "Managed Postgres with a free tier, session-pooler URL that survives long-lived FastAPI workers."],
            ["Object storage", "Supabase Storage", "supabase-py 2.10.0",
             "Holds the original PDF for audit. Gracefully degrades if not configured — extraction still works."],
            ["Auth", "Custom JWT (HS256)", "python-jose 3.3.0",
             "Bcrypt for hashing, JWT for stateless auth. Simpler than wiring Supabase Auth for a single-day demo and gives us full control over claims (role)."],
            ["Password hashing", "passlib + bcrypt", "1.7.4",
             "Industry-standard. We never log raw passwords."],
            ["AI provider", "Google Gemini via google-genai", "0.3.0",
             "Best price/performance on the free tier. 2.5-pro for ranking (reasoning-heavy), 2.5-flash for inference and JD distillation."],
            ["PDF extraction", "pypdf + pdfplumber", "5.1.0 / 0.11.4",
             "pypdf for fast text extraction; pdfplumber kept as a fallback dependency for tricky layouts."],
            ["SSE streaming", "Native FastAPI StreamingResponse", "—",
             "Plain text/event-stream — no extra library needed. The frontend uses fetch + ReadableStream."],
        ],
        col_widths=[1.4, 1.7, 1.0, 2.4],
    )

    add_heading(doc, "Architecture Decisions", level=2)

    add_para(
        doc,
        "JWT over Supabase Auth. Supabase Auth would have been one less moving "
        "piece, but it would also bind the auth UX to the Supabase JS SDK. We need "
        "a stateless JWT with a custom claim (role) for our edge middleware to "
        "decode and route by, and we need bcrypt hashes we control. Our auth fits "
        "in ~120 lines (security.py + auth/service.py) and the trade-off — we "
        "don't get social login or password recovery for free — is acceptable for "
        "Phase 1.",
    )
    add_para(
        doc,
        "Single-call AI ranking over embedding + rerank. With 15–30 approved "
        "profiles, an embedding prefilter introduces a recall risk (the right "
        "candidate gets cut at retrieval) for no real latency win — the LLM call "
        "is the dominant cost. A single ranking call also produces dramatically "
        "better reasoning, which is what the demo trades on. At >1k profiles we "
        "would absolutely add an embedding prefilter (Supabase pgvector or "
        "OpenAI/Cohere embeddings) and reuse the same prompt as the reranker. "
        "Phase 1 is the small-N regime where one strong prompt wins.",
    )
    add_para(
        doc,
        "Why Gemini, and not Anthropic Claude. The hackathon brief mentioned "
        "both Anthropic env vars and Google Gemini in the tech-stack section. "
        "We picked Gemini for three pragmatic reasons: a free-tier daily quota "
        "we could exercise during development; 2.5-pro's reasoning quality on "
        "structured-JSON tasks is competitive with Claude for our use case; "
        "and the google-genai SDK's streaming generator works cleanly with "
        "our SSE pipeline. The whole AI integration is one file "
        "(modules/ai/client.py) — swapping to Anthropic SDK is roughly 30 lines "
        "if Phase 2 prioritises Claude.",
    )
    add_para(
        doc,
        "Model tier split. GEMINI_MODEL_SHOWCASE (default gemini-2.5-pro) is "
        "used for resume extraction and search ranking — the two places where "
        "reasoning quality is visible to the user. GEMINI_MODEL_LIGHT (default "
        "gemini-2.5-flash) handles related-skill inference and JD distillation — "
        "deterministic rewriting tasks where flash's lower latency is the win. "
        "Both are settable in .env so the model choice is configuration, not "
        "code.",
    )

    add_page_break(doc)

    # ─────────────────────────────────────────────────────────────
    # 5. SYSTEM ARCHITECTURE
    # ─────────────────────────────────────────────────────────────
    add_heading(doc, "5. System Architecture", level=1)

    add_heading(doc, "5.1 High-Level Architecture Diagram", level=2)
    add_code_block(
        doc,
        """\
+---------------------------+        +-------------------------------------+
|  Browser                  |        |  FastAPI (Render / Docker)          |
|  Next.js 14 (Vercel)      |        |                                     |
|  - 2 magic screens        | <----> |  /auth     (login, signup, /me)     |
|  - SSE result stream      |  HTTPS |  /resumes  (upload + AI extract)    |
|  - Edge middleware        |  JWT   |  /search   (NL + JD + Team SSE)     |
|    decodes JWT, routes    |        |  /review   (HR queue, approve)      |
|    by role                |        |  /employees (CRUD + directory)      |
+---------------------------+        |                                     |
            ^                        |  Async SQLAlchemy 2.0               |
            |                        |  google-genai client                |
            |  multipart PDF         +-----------+------+------------------+
            |                                    |      |
            v                                    v      v
   +-----------------------+      +-------------------+ +----------------+
   |  Supabase Storage     |      |  Supabase         | |  Google Gemini |
   |  (resumes/)           |      |  Postgres         | |  2.5-pro       |
   |  original PDFs        |      |  asyncpg pooler   | |  2.5-flash     |
   +-----------------------+      +-------------------+ +----------------+

Auth: JWT in `Authorization: Bearer ...` header (API).  Edge middleware
reads a non-HttpOnly `skillshub_token` cookie set in parallel by the
frontend so it can route by role before the page renders.
""",
    )

    add_heading(doc, "5.2 Backend Module Structure", level=2)
    add_para(doc, "Every module owns its router, service, schemas, and (where applicable) models. Cross-cutting helpers live in app/common and app/core.")
    add_code_block(
        doc,
        """\
backend/app/
  core/
    config.py         pydantic-settings; validated at startup
    database.py       async engine + AsyncSessionLocal + get_db dep
    security.py       JWT encode/decode, bcrypt hash/verify
    exceptions.py     AppError hierarchy (4xx + UpstreamAIError 502)

  common/
    dependencies.py   CurrentUser, DbSession, require_role
    decorators.py     require_admin = require_role(UserRole.ADMIN)
    middleware.py     RequestIdMiddleware (request id + access log)
    filters.py        Exception handlers (AppError -> JSON envelope)
    rate_limit.py     In-memory sliding window for /auth/* endpoints
    utils.py          safe_json_loads (strips code fences), initials

  modules/
    auth/             register/employee, register/hr, login, /auth/me
    users/            User model, /users/me
    employees/        Employee/Skill/Project models + directory CRUD
    resumes/          POST /resumes/upload — the AI extraction pipeline
    review/           HR review queue: list, get, approve, reject
    search/           POST /search (SSE), /search/jd, /search/team
    ai/
      client.py       Singleton google-genai client
      service.py      extract_resume, infer_related_skills,
                      stream_search_results, build_team, distill_jd_to_query
      schemas.py      Pydantic models for AI input/output
      prompts/        Versioned prompts: extraction_v1, inference_v1,
                      search_v1, jd_distill_v1

  alembic/versions/
    0001_initial.py   Full Phase 1 schema in one migration

  scripts/
    seed_database.py  Idempotent reseed: 2 demo users + 15 profiles
    extract_one.py    CLI: run the extraction prompt on a single resume
    search_one.py     CLI: run one search query and print the ranking
""",
    )

    add_heading(doc, "5.3 Frontend Route Structure", level=2)
    add_para(doc, "Route groups in parentheses don't appear in URLs — they just bind a shared layout (and role check) to a set of pages.")
    add_code_block(
        doc,
        """\
frontend/app/
  layout.tsx              Root layout: Providers (QueryClient, Toaster)
  page.tsx                / -> redirects to /login
  middleware.ts           Edge middleware: JWT decode + role-aware redirects
                          and `?notice=role_mismatch_*` query param hand-off

  (auth)/
    login/page.tsx        Single page with sign-in tab + signup tab + role
                          switcher + 2 one-click demo buttons

  (hr)/                   ADMIN-only. Layout shows toast on role_mismatch.
    layout.tsx            AppShell role="ADMIN" + NoticeHandler
    search/page.tsx       3 tabs: Search | From JD | Build Team
    directory/page.tsx    Filterable employee grid
    review/page.tsx       Pending review queue
    review/[id]/page.tsx  Single-item editor + approve/reject

  (employee)/             USER-only. Layout shows toast on role_mismatch.
    layout.tsx            AppShell role="USER" + NoticeHandler
    upload/page.tsx       The resume-upload magic screen
    profile/page.tsx      Self-view + edit

frontend/components/
  ui/                     shadcn primitives (button, card, dialog, etc.)
  shared/                 app-shell, avatar, score-badge, skill-pill,
                          streaming-indicator, employee-sheet
  features/
    auth/                 sign-in, signup-employee, signup-hr,
                          password-strength-indicator, demo-access-card
    resume-upload/        dropzone, extraction-progress, profile-editor
    search/               search-input, result-card

frontend/hooks/           use-auth, use-search, use-resume-upload,
                          use-employees, use-review-queue
frontend/lib/             api.ts (axios), auth.ts (localStorage + cookies),
                          validations.ts (zod), utils.ts
""",
    )

    add_heading(doc, "5.4 Data Model", level=2)

    add_para(doc, "User", bold=True)
    add_table(
        doc,
        ["Field", "Type", "Constraints", "Purpose"],
        [
            ["id", "UUID", "PK, default uuid4", "Subject of every JWT we issue."],
            ["name", "String(50)", "Required", "Display name on the app shell + result cards."],
            ["email", "String(100)", "Unique, indexed", "Login identity. Lowercased on write."],
            ["password", "String", "Required", "Bcrypt hash (never the plaintext)."],
            ["role", "Enum(ADMIN | USER)", "Default USER", "Drives RBAC at every protected endpoint."],
            ["created_at, updated_at", "DateTime(tz)", "Server defaults", "Audit trail."],
        ],
        col_widths=[1.4, 1.4, 1.4, 2.5],
    )
    add_para(
        doc,
        "Relationships: a User has at most one Employee (back-ref \"employee\"). "
        "When a USER signs up, a stub Employee row is created so the resume-upload "
        "flow has somewhere to land. ADMIN users do not have an Employee row.",
    )

    add_para(doc, "Employee", bold=True)
    add_table(
        doc,
        ["Field", "Type", "Constraints", "Purpose"],
        [
            ["id", "UUID", "PK", "Stable identifier used in URLs and JWT claims."],
            ["user_id", "UUID", "FK→users.id, nullable, unique", "Null for HR-uploaded profiles; unique so one user maps to one employee."],
            ["full_name", "String(100)", "Required", "Source-of-truth name (may differ from User.name)."],
            ["headline", "String(200)", "Nullable", "One-liner like \"Senior Frontend Engineer\"."],
            ["location", "String(100)", "Indexed, nullable", "Single city; \"Pune\", \"Bangalore\", \"Remote\"."],
            ["years_experience", "Numeric(4,1)", "Nullable", "Overall career span."],
            ["allocation_status", "Enum", "Indexed; UNALLOCATED|PARTIAL|ALLOCATED", "Hard signal for availability searches."],
            ["last_project_end_date", "Date", "Nullable", "Used by the temporal search rules."],
            ["bio", "Text", "Nullable", "Free-form summary; editable in profile."],
            ["resume_url", "Text", "Nullable", "Supabase Storage public URL."],
            ["raw_extracted_json", "JSONB", "Nullable", "The full Gemini extraction blob — preserved for audit."],
            ["status", "Enum", "Indexed; PENDING_REVIEW|APPROVED|REJECTED", "Gatekeeper for search visibility."],
            ["created_at, updated_at", "DateTime(tz)", "Server defaults", "Audit trail."],
        ],
        col_widths=[1.5, 1.3, 1.5, 2.4],
    )
    add_para(
        doc,
        "Relationships: skills (1-N, cascade delete), projects (1-N, cascade), "
        "review_items (1-N, cascade). Only Employee rows with status=APPROVED are "
        "ever passed to the search prompt (see EmployeeService.list_approved_candidates).",
    )

    add_para(doc, "Skill", bold=True)
    add_table(
        doc,
        ["Field", "Type", "Constraints", "Purpose"],
        [
            ["id", "UUID", "PK", "—"],
            ["employee_id", "UUID", "FK, cascade, indexed", "Owning employee."],
            ["name", "String(80)", "Indexed; unique per employee", "Normalised skill name like \"Spring Boot\"."],
            ["category", "Enum", "LANGUAGE|FRAMEWORK|PLATFORM|TOOL|DOMAIN", "Filter facet and demo grouping."],
            ["proficiency", "Enum", "NOVICE|INTERMEDIATE|EXPERT", "Inferred from verb choice; see extraction prompt rules."],
            ["years_experience", "Numeric(4,1)", "Nullable", "Per-skill, summed from projects where it appears."],
            ["is_inferred", "Boolean", "Default false", "True for skills the inference call added."],
            ["inference_confidence", "Numeric(3,2)", "Nullable, 0–1", "Only set for inferred skills."],
            ["inference_reason", "Text", "Nullable", "Why the model thought this skill applies."],
            ["evidence", "Text", "Nullable", "Quote ≤180 chars from the resume that justifies the skill."],
        ],
        col_widths=[1.5, 1.3, 1.7, 2.2],
    )

    add_para(doc, "Project", bold=True)
    add_table(
        doc,
        ["Field", "Type", "Constraints", "Purpose"],
        [
            ["id", "UUID", "PK", "—"],
            ["employee_id", "UUID", "FK, cascade, indexed", "Owning employee."],
            ["title", "String(200)", "Required", "Used in evidence_snippets on result cards."],
            ["description", "Text", "Nullable", "Short paragraph; passed to the search prompt for ranking context."],
            ["role", "String(100)", "Nullable", "\"Tech Lead\", \"Senior Engineer\", etc."],
            ["domain", "String(80)", "Indexed, nullable", "\"Fintech\", \"Healthcare\", \"E-commerce\"."],
            ["start_date, end_date", "Date", "Nullable", "Drives last_project_end_date and temporal search."],
            ["tech_stack", "JSONB", "Nullable", "List of skill names. Uses same vocabulary as the skills table."],
        ],
        col_widths=[1.5, 1.3, 1.6, 2.3],
    )

    add_para(doc, "ReviewQueueItem", bold=True)
    add_table(
        doc,
        ["Field", "Type", "Constraints", "Purpose"],
        [
            ["id", "UUID", "PK", "—"],
            ["employee_id", "UUID", "FK, cascade, indexed", "Profile under review."],
            ["submitted_by_user_id", "UUID", "FK, nullable, SET NULL on delete", "Whose upload triggered this."],
            ["status", "Enum", "Indexed; PENDING|APPROVED|REJECTED|EDITED_AND_APPROVED", "Workflow state."],
            ["reviewer_id", "UUID", "FK, nullable, SET NULL", "Which ADMIN took the decision."],
            ["reviewer_notes", "Text", "Nullable", "Free-form notes saved with the decision."],
            ["created_at, reviewed_at", "DateTime(tz)", "Server defaults", "Audit trail."],
        ],
        col_widths=[1.6, 1.3, 1.8, 2.1],
    )

    add_para(doc, "SearchQueryLog", bold=True)
    add_table(
        doc,
        ["Field", "Type", "Constraints", "Purpose"],
        [
            ["id", "UUID", "PK", "—"],
            ["user_id", "UUID", "FK, nullable, SET NULL", "Who ran the query (HR)."],
            ["query_text", "Text", "Required", "Verbatim natural-language query."],
            ["result_count", "Integer", "Default 0", "How many results streamed (for quality analysis)."],
            ["top_match_score", "Integer", "Nullable", "Highest match_score returned — a rough quality signal."],
            ["created_at", "DateTime(tz)", "Server default", "Audit trail."],
        ],
        col_widths=[1.5, 1.3, 1.6, 2.3],
    )
    add_para(
        doc,
        "Every successful or partial search writes one row here. This is the "
        "audit log judges asked about, and it's what would feed a future "
        "prompt-quality dashboard.",
    )

    add_page_break(doc)

    # ─────────────────────────────────────────────────────────────
    # 6. END-TO-END WORKING FLOWS
    # ─────────────────────────────────────────────────────────────
    add_heading(doc, "6. End-to-End Working Flows", level=1)
    add_para(
        doc,
        "Each flow below walks through the code path a new team member would "
        "follow with a debugger. File and function names are exact.",
    )

    add_heading(doc, "6.1 New Employee Signs Up & Uploads Resume", level=2)
    add_para(doc, "One sentence: a fresh USER account is created, a stub Employee row is reserved, and a single PDF upload becomes a structured, AI-extracted profile sitting in the review queue.")
    add_para(doc, "Triggering action: visitor opens /login, switches to the Create Account tab, picks \"I'm an Employee\", submits the form. Then drags a PDF onto the dropzone at /upload.")
    add_para(doc, "Step-by-step:", bold=True)
    add_numbered(doc, "frontend/app/(auth)/login/page.tsx renders the SignupEmployeeForm. On submit, react-hook-form + Zod validate locally (≥8 chars, ≥1 letter, ≥1 digit, password === confirm_password).")
    add_numbered(doc, "hooks/use-auth.ts:useSignupEmployee POSTs to /auth/register/employee. The auth_rate_limit dependency in backend/app/modules/auth/router.py caps this to 10 requests/min/IP.")
    add_numbered(doc, "AuthService.register_employee (modules/auth/service.py) checks settings.ALLOW_EMPLOYEE_SIGNUP, creates a User (UserRole.USER), then creates a stub Employee with status=PENDING_REVIEW and allocation_status=UNALLOCATED, then commits.")
    add_numbered(doc, "The response carries access_token, user, employee_id, and next_action='upload_resume'. lib/auth.ts:setSession writes both localStorage and the skillshub_token + skillshub_role cookies for the edge middleware.")
    add_numbered(doc, "useSignupEmployee.onSuccess routes the browser to /upload. middleware.ts decodes the JWT and confirms USER role for the (employee) route group.")
    add_numbered(doc, "app/(employee)/upload/page.tsx mounts. The user drops a PDF. components/features/resume-upload/dropzone.tsx accepts only application/pdf and ≤5 MB.")
    add_numbered(doc, "hooks/use-resume-upload.ts:useResumeUpload posts multipart/form-data to /resumes/upload. The page flips to the \"processing\" phase and shows components/features/resume-upload/extraction-progress.tsx — a 4-step animated checklist.")
    add_numbered(doc, "Backend: ResumeService.upload_and_extract (modules/resumes/service.py) validates the PDF: ≥200 bytes, ≤5 MB, ends in .pdf.")
    add_numbered(doc, "_pdf_to_text uses pypdf to extract text. If <80 chars, it raises BadRequestError(\"PDF text was too short to parse — is this a scanned image?\").")
    add_numbered(doc, "_upload_to_supabase pushes the raw PDF to Supabase Storage bucket \"resumes/\" (best-effort — if unconfigured, returns None and the rest still works).")
    add_numbered(doc, "ai_service.extract_resume (modules/ai/service.py) builds the prompt from prompts/extraction_v1.py and calls Gemini (settings.GEMINI_MODEL_SHOWCASE). Returns an ExtractedProfile pydantic object.")
    add_numbered(doc, "ai_service.infer_related_skills runs a second Gemini call with prompts/inference_v1.py — at most 5 inferences, confidence ≥0.7, deduped against the explicit skill list. Errors here degrade gracefully to [].")
    add_numbered(doc, "ResumeService._persist either updates the existing stub Employee (for USER role) or creates a fresh one (for HR uploading on behalf). Skills are replaced en bloc; projects are replaced; latest project end_date is computed and stored in last_project_end_date.")
    add_numbered(doc, "A ReviewQueueItem is created with status=PENDING. db.commit().")
    add_numbered(doc, "The endpoint returns ResumeUploadResponse: the full employee, the raw ExtractedProfile, the inferred_skills list, and the resume_url. Total wall time: ~6–10 s.")
    add_numbered(doc, "Frontend: upload page swaps to the \"ready\" phase. ProfileEditor renders editable skill pills (with confidence % tooltip on inferred skills) and project cards. Toast: \"Resume parsed. Review and send for approval.\"")
    add_numbered(doc, "User clicks Send for Review. update + replace_skills + replace_projects are PATCHed/PUT'd in sequence. The profile sits in PENDING_REVIEW until HR approves it.")
    add_para(doc, "Where AI is involved:", bold=True)
    add_bullet(doc, "Step 11 — resume extraction (gemini-2.5-pro by default, extraction_v1.py).")
    add_bullet(doc, "Step 12 — related-skill inference (gemini-2.5-flash, inference_v1.py).")
    add_para(doc, "End state:", bold=True)
    add_bullet(doc, "users row exists with role=USER and a bcrypt hash.")
    add_bullet(doc, "employees row in status=PENDING_REVIEW with the full extracted name/headline/location/years.")
    add_bullet(doc, "skills rows: explicit ones with is_inferred=false + evidence; inferred ones with is_inferred=true + inference_confidence + inference_reason.")
    add_bullet(doc, "projects rows with start_date, end_date, tech_stack.")
    add_bullet(doc, "One review_queue_items row with status=PENDING.")
    add_bullet(doc, "Optional: resume_url pointing at Supabase Storage.")
    add_para(doc, "Common failure modes:", bold=True)
    add_bullet(doc, "Scanned PDF (image-only) → BadRequestError surfaced as toast \"PDF text was too short to parse — is this a scanned image?\".")
    add_bullet(doc, "Gemini quota exhausted → UpstreamAIError with the friendly message wired in _humanize_gemini_error (\"Daily AI quota reached for the configured Gemini model. …\").")
    add_bullet(doc, "Inference call fails → swallowed; profile saves with explicit skills only. No user-visible error.")
    add_bullet(doc, "Supabase Storage misconfigured → resume_url is null; no PDF download link on the review page, but everything else works.")

    add_heading(doc, "6.2 HR Reviews and Approves an Extracted Profile", level=2)
    add_para(doc, "One sentence: HR opens the queue, reviews an extracted profile against the original PDF, optionally edits any field, and clicks Approve — which flips status to APPROVED and makes the candidate searchable.")
    add_para(doc, "Triggering action: HR clicks Review in the app shell (or follows a deep link).")
    add_para(doc, "Step-by-step:", bold=True)
    add_numbered(doc, "Edge middleware confirms the JWT decodes to role=ADMIN and lets /review render.")
    add_numbered(doc, "app/(hr)/review/page.tsx mounts. useReviewQueue hits GET /review/queue. require_admin dependency on the entire review router ensures only ADMIN reaches the handler.")
    add_numbered(doc, "ReviewService.list_pending (modules/review/service.py) returns all ReviewQueueItem rows where status=PENDING, ordered by created_at ascending.")
    add_numbered(doc, "User clicks Review on a row → /review/[id] page. GET /review/{item_id} returns ReviewQueueItemWithEmployee — the review row plus a fully-loaded EmployeeResponse with skills and projects.")
    add_numbered(doc, "The page renders the ProfileEditor (the same one used on /upload). HR can edit every field: name, headline, location, allocation_status, skills (rename, delete, change proficiency/category), projects.")
    add_numbered(doc, "If HR edits anything, the same PATCH /employees/{id} + PUT /employees/{id}/skills + PUT /employees/{id}/projects sequence runs. EmployeeService.ensure_can_edit allows ADMIN to edit any profile (USERs can only edit their own).")
    add_numbered(doc, "HR clicks Approve. POST /review/{id}/approve with optional notes. ReviewService.approve sets the review item to APPROVED, stamps reviewer_id + reviewed_at, then calls EmployeeService.set_status(employee_id, ProfileStatus.APPROVED). db.commit().")
    add_numbered(doc, "From this point on, the employee appears in /search results (search service's list_approved_candidates only returns status=APPROVED).")
    add_para(doc, "Where AI is involved: none — this flow is purely human review.", bold=False)
    add_para(doc, "End state:", bold=True)
    add_bullet(doc, "review_queue_items row: status=APPROVED, reviewer_id set, reviewed_at set, reviewer_notes optional.")
    add_bullet(doc, "employees row: status=APPROVED. The profile is now in the search corpus.")
    add_para(doc, "Common failure modes:", bold=True)
    add_bullet(doc, "Reviewer tries to approve an already-resolved item → ConflictError 409, frontend toasts the message.")
    add_bullet(doc, "Non-ADMIN attempts /review → backend returns 403 from require_admin; middleware would have already redirected to /search with notice=role_mismatch_user.")

    add_heading(doc, "6.3 HR Runs a Natural-Language Search", level=2)
    add_para(doc, "One sentence: HR types a real question, the backend ships every approved profile plus a temporal-aware prompt to Gemini, and the frontend streams ranked candidate cards as Gemini emits each JSON object.")
    add_para(doc, "Triggering action: HR types in the search input on /search and presses Enter.")
    add_para(doc, "Step-by-step:", bold=True)
    add_numbered(doc, "frontend/hooks/use-search.ts:useSearch.run posts {query, limit} to /search using fetch + AbortController (EventSource doesn't support custom Authorization headers).")
    add_numbered(doc, "Backend: search/router.py:search_stream gates on require_admin and returns a StreamingResponse with media_type=text/event-stream and headers X-Accel-Buffering: no (disables Nginx buffering).")
    add_numbered(doc, "SearchService.stream (modules/search/service.py) starts: today = date.today(); _augment_with_recency injects days_since_last_project_end into every candidate; _temporal_context yields today/last_month_start/last_quarter_start.")
    add_numbered(doc, "EmployeeService.list_approved_candidates loads every APPROVED Employee with skills + projects eagerly loaded (selectinload). EmployeeService.to_search_candidate compacts each row to the minimal JSON the prompt needs.")
    add_numbered(doc, "ai_service.stream_search_results is called with (query, candidates, limit, temporal_context). It builds the full prompt from prompts/search_v1.py and calls google-genai client.models.generate_content_stream against settings.GEMINI_MODEL_SHOWCASE.")
    add_numbered(doc, "Producer/consumer pattern: the SDK is sync, so the call runs in a worker thread feeding an asyncio.Queue. The buffer accumulates chunks; _iter_complete_objects tracks brace depth + string state to extract every complete top-level JSON object.")
    add_numbered(doc, "Each complete object is validated against SearchResult (pydantic) and yielded. The router formats it as event: result\\ndata: {json}\\n\\n.")
    add_numbered(doc, "Frontend: use-search.ts reads ReadableStream chunks, splits on \\n\\n, parses event: result frames and pushes them into the results state. Each ResultCard slides in with framer-motion.")
    add_numbered(doc, "When the stream ends (or hits an error), SearchService writes one SearchQueryLog row with query_text, result_count, top_match_score, user_id. db.commit().")
    add_numbered(doc, "Final event is event: done\\ndata: {\"done\": true, \"count\": N}\\n\\n. Frontend flips phase to \"done\" and shows the result count.")
    add_para(doc, "Where AI is involved: step 5 — the showcase Gemini call (search_v1.py).", bold=False)
    add_para(doc, "End state:", bold=True)
    add_bullet(doc, "Up to `limit` SearchResult cards rendered, each with: match_score (0–100), reason (2–3 sentences), matched_skill_names (highlighted on the card), evidence_snippets (visible in the \"Why?\" tooltip).")
    add_bullet(doc, "One search_query_logs row written for audit/analysis.")
    add_para(doc, "Common failure modes:", bold=True)
    add_bullet(doc, "Empty corpus (no APPROVED profiles) → emits event: done with message: \"No approved profiles to search.\".")
    add_bullet(doc, "Gemini quota / 5xx → emits event: error\\ndata: {message: <_humanize_gemini_error>}. Frontend shows the humanised message inline (\"Daily AI quota reached…\", \"AI service is temporarily overloaded…\").")
    add_bullet(doc, "Malformed JSON in a chunk → that one chunk is skipped, _iter_complete_objects continues with the next.")
    add_bullet(doc, "Network drop mid-stream → AbortController on the client cancels the fetch; backend stream terminates cleanly.")

    add_heading(doc, "6.4 HR Builds a Team (and the JD-to-Shortlist Variant)", level=2)
    add_para(doc, "Both Team Builder and JD-to-Shortlist are implemented in Phase 1, sharing the same approved-candidate corpus as POST /search.")
    add_para(doc, "Team Builder.", bold=True)
    add_numbered(doc, "/search page → Build Team tab. HR types a brief like \"4-person team for a 3-month healthcare app, needs mobile + backend + DevOps\" and picks a team size (2–8).")
    add_numbered(doc, "hooks/use-search.ts:useTeamBuilder posts to /search/team. Single, non-streaming call.")
    add_numbered(doc, "SearchService.build_team passes the brief to ai_service.build_team, which uses prompts/search_v1.py:build_team_prompt (a sibling of the search prompt).")
    add_numbered(doc, "Gemini returns TeamBuildResult: a team (exactly team_size members, each with role_on_team + why_picked), a rationale paragraph, and alternates.")
    add_numbered(doc, "The frontend renders the rationale card on top, then a 2-column grid of team-member cards, and an alternates list below. Each member card is clickable into the EmployeeSheet.")
    add_para(doc, "JD-to-Shortlist (paste-a-JD search).", bold=True)
    add_numbered(doc, "/search page → From JD tab. HR pastes a verbose job description (≥30 chars, ≤8 000).")
    add_numbered(doc, "hooks/use-search.ts:useSearch.runFromJD posts to /search/jd as SSE.")
    add_numbered(doc, "Stage 1: SearchService.stream_from_jd calls ai_service.distill_jd_to_query (prompts/jd_distill_v1.py, light model) to compress the JD into a single ≤200-char hiring query.")
    add_numbered(doc, "First SSE event is event: query\\ndata: {generated_query: \"...\"}. Frontend renders a primary-tinted pill so HR can see what the AI heard.")
    add_numbered(doc, "Stage 2: stream_from_jd then delegates to SearchService.stream(generated_query, ...) — same ranker, same streaming pattern, same audit log.")
    add_para(doc, "Common failure modes:", bold=True)
    add_bullet(doc, "Distillation fails (empty output, quota) → event: error then event: done with count: 0. Frontend shows the error inline.")
    add_bullet(doc, "Team builder gets an unrealistic brief (no DevOps people in the corpus) → the rationale paragraph will say so and the team list will include the closest fit + alternates.")

    add_page_break(doc)

    # ─────────────────────────────────────────────────────────────
    # 7. AI INTEGRATION DEEP DIVE
    # ─────────────────────────────────────────────────────────────
    add_heading(doc, "7. AI Integration Deep Dive", level=1)

    add_heading(doc, "7.1 Why AI, and Why Gemini", level=2)
    add_para(
        doc,
        "Both centerpiece features are unsolvable with classical NLP at hackathon "
        "scale. Resume parsing with regex fails on the variety real resumes "
        "exhibit — Indian and US date formats, MM-DD-YYYY vs YYYY-MM, \"Tech "
        "Lead\" vs \"Engineering Lead\", project blurbs that bury the tech "
        "stack in prose. A modern LLM with a strict JSON schema solves it in "
        "one call.",
    )
    add_para(
        doc,
        "Search ranking with full-text or BM25 has the inverse problem: it "
        "matches strings, not meaning. \"Payment gateway\" must match Razorpay, "
        "Stripe, PayU, Paytm, Square, Braintree. \"Real-time\" must match "
        "Socket.IO, WebSocket, WebRTC, SSE. \"ML\" must match PyTorch, "
        "TensorFlow, scikit-learn. Encoding all of those equivalences "
        "manually is the project. An LLM, with a few-shot prompt and a "
        "structured-output schema, just knows.",
    )
    add_para(
        doc,
        "We picked Google Gemini for the practical reasons in §4: a usable "
        "free-tier daily quota for development, 2.5-pro's reasoning quality "
        "on structured JSON, and a clean sync-streaming SDK that we wrap in "
        "an asyncio.Queue. The brief mentioned Anthropic env vars alongside "
        "Gemini in its tech-stack section; we resolved the ambiguity in favour "
        "of the simpler path and isolated the choice to one file "
        "(modules/ai/client.py).",
    )

    add_heading(doc, "7.2 The Four AI Calls", level=2)

    add_para(doc, "Call 1: Resume Extraction (extraction_v1.py)", bold=True)
    add_para(doc, "Model: settings.GEMINI_MODEL_SHOWCASE (default gemini-2.5-pro). Temperature: 0.1.")
    add_para(doc, "Purpose: turn ~1–3 pages of resume text into a strict ExtractedProfile JSON with skills (category + proficiency + evidence quote), projects (with dates + tech_stack), and headline info.")
    add_para(doc, "System prompt + rules (verbatim from extraction_v1.py):", italic=True)
    add_code_block(
        doc,
        """\
You are a precise resume parser for an enterprise HR system.
Output ONLY valid JSON matching the schema below. No prose. No markdown fences.
No commentary. If a value is unknown, return null. NEVER invent facts.

EXTRACTION RULES (follow strictly):
1. Proficiency is inferred from verb choice and context:
   - "architected", "led", "owned", "designed end-to-end" -> EXPERT
   - "built", "developed", "contributed to", "implemented" -> INTERMEDIATE
   - "used", "familiar with", "exposed to", "learned"      -> NOVICE
2. Compute years_experience per skill by summing the date ranges of projects
   in which that skill appears. If dates are missing, estimate conservatively
   (round DOWN).
3. Every skill MUST include a short `evidence` quote (<=180 chars) taken from
   the resume — the exact phrase that justified picking it.
4. Categorise skills:
   - LANGUAGE: Java, Python, TypeScript, Go, etc.
   - FRAMEWORK: React, Spring Boot, FastAPI, Next.js, etc.
   - PLATFORM: AWS, GCP, Kubernetes, Docker, Vercel, Supabase, etc.
   - TOOL: Git, Postman, Jira, Figma, etc.
   - DOMAIN: Fintech, Healthcare, E-commerce, Logistics, etc.
5. Project tech_stack must use the SAME normalised skill names you used in
   `skills`.
6. Location should be a single city name where possible ("Pune", "Bangalore",
   "Remote").
7. years_experience at profile level = the overall career span, not the sum
   of skills.
""",
    )
    add_para(doc, "Few-shot example output (excerpt):", italic=True)
    add_code_block(
        doc,
        """\
{
  "full_name": "Priya Sharma",
  "headline": "Senior Backend Engineer",
  "location": "Bangalore",
  "years_experience": 6,
  "skills": [
    {"name": "Java", "category": "LANGUAGE", "proficiency": "EXPERT",
     "years_experience": 6,
     "evidence": "6 years building scalable APIs ... Spring Boot microservices"},
    {"name": "Razorpay", "category": "PLATFORM", "proficiency": "EXPERT",
     "years_experience": 3,
     "evidence": "Architected the Razorpay integration serving 2M+ monthly transactions"},
    ...
  ],
  "projects": [
    {"title": "Razorpay Payment Integration", "role": "Tech Lead",
     "domain": "Fintech", "start_date": "2022-01", "end_date": null,
     "tech_stack": ["Java", "Spring Boot", "PostgreSQL", "Razorpay"]}
  ]
}
""",
    )
    add_para(doc, "Why this design matters: the evidence field is the anti-hallucination guardrail. The model is forbidden from inventing skills — it must quote the resume. Categories normalise the taxonomy so the directory filter works. The verb-based proficiency rule is deterministic enough to be reproducible across runs.")

    add_para(doc, "Call 2: Related-Skill Inference (inference_v1.py)", bold=True)
    add_para(doc, "Model: settings.GEMINI_MODEL_LIGHT (default gemini-2.5-flash). Temperature: 0.3.")
    add_para(doc, "Purpose: given the explicit skill list extracted in Call 1, infer up to 5 high-confidence (≥0.7) related skills with a one-line reason. These appear in the UI in a separate \"Inferred Skills\" section so HR can tell them apart.")
    add_para(doc, "Full system prompt (verbatim):", italic=True)
    add_code_block(
        doc,
        """\
You infer skills a developer almost certainly has, given the skills they have
explicitly listed on their resume. Return ONLY high-confidence inferences
(confidence >= 0.7). Output STRICT JSON. No prose, no markdown fences.

Rules:
- Maximum 5 inferences. Quality > quantity.
- Do NOT duplicate any skill that already appears in the input list.
- Each inference must be supported by a clear technical reason (e.g.
  dependency, shared toolchain, ecosystem co-membership).
- Confidence is a float in [0.7, 1.0]. Use higher values only when the
  relationship is essentially mechanical (Next.js -> React) and lower when
  it's strong-but-soft (Java -> JUnit).
- Proficiency must NOT exceed the proficiency of the strongest source skill.
""",
    )
    add_para(doc, "Sample output:", italic=True)
    add_code_block(
        doc,
        """\
[
  {"name": "React", "category": "FRAMEWORK", "proficiency": "EXPERT",
   "confidence": 0.98,
   "reason": "Next.js is built on top of React; 4 years of Next.js implies
              expert React knowledge."},
  {"name": "JavaScript", "category": "LANGUAGE", "proficiency": "EXPERT",
   "confidence": 0.97,
   "reason": "TypeScript is a superset of JavaScript and Next.js compiles to JS."}
]
""",
    )
    add_para(doc, "Why this design matters: dedup is enforced client-side (in modules/ai/service.py:infer_related_skills) against the explicit names — the model is told to dedupe, but we don't trust it to. Confidence floor (0.7) keeps low-signal noise out. The reason field is what shows up in the UI tooltip.")

    add_para(doc, "Call 3: Semantic Search Ranking (search_v1.py)", bold=True)
    add_para(doc, "Model: settings.GEMINI_MODEL_SHOWCASE (default gemini-2.5-pro). Temperature: 0.2.")
    add_para(doc, "Purpose: rank up to N (1–10) candidates against a natural-language hiring query. Stream JSON objects so the frontend can render cards as they arrive.")
    add_para(doc, "Key rules from the system prompt (excerpt):", italic=True)
    add_code_block(
        doc,
        """\
RANKING RULES:
- Understand the query SEMANTICALLY:
   * "payment gateway" matches Razorpay, Stripe, PayU, Paytm, Square, Braintree.
   * "real-time" matches Socket.IO, WebSocket, WebRTC, SSE, Firebase RTDB.
   * "ML" matches PyTorch, TensorFlow, scikit-learn, Hugging Face, LangChain.

- Interpret AVAILABILITY and TIME constraints using the TEMPORAL CONTEXT
  block below, plus the per-candidate fields `allocation_status`,
  `last_project_end_date`, and `days_since_last_project_end`. NEVER invent
  dates — use the values given.
   * "available now" -> allocation_status == UNALLOCATED.
   * "recently shipped", "in the last quarter" ->
       0 <= days_since_last_project_end <= 90.
   * "haven't been on a new project in the last quarter", "stale" ->
       days_since_last_project_end > 90 OR allocation_status == UNALLOCATED.

- match_score in [0, 100]:
   * 90-100: hits every hard requirement with strong evidence + seniority fit.
   * 75-89:  hits every hard requirement, weaker on one nice-to-have.
   * 60-74:  misses one hard requirement OR strong on skills but wrong location/seniority.
   * <60:    significant gap.

- `reason` must cite SPECIFIC EVIDENCE — years, project names, roles, domain.
  NEVER vague phrases like "good fit" or "strong candidate".
- `matched_skill_names` must be exact strings copied from the candidate's
  skills array, so the UI can highlight them.
- `evidence_snippets` must quote 1-3 short fragments from the candidate's
  projects or headline that justify the match.
""",
    )
    add_para(doc, "Dynamic injection: build_prompt also emits a TEMPORAL CONTEXT section computed at request time (today, last_month_start, last_quarter_start) and each candidate carries a precomputed days_since_last_project_end value. This is what makes \"hasn't shipped in the last quarter\" reliably correct.")
    add_para(doc, "Sample streamed element:", italic=True)
    add_code_block(
        doc,
        """\
{
  "employee_id": "b532b69e-d640-4c2e-a1ae-9880b6630c2e",
  "name": "Vikram Kulkarni",
  "headline": "Senior Frontend Engineer",
  "location": "Pune",
  "allocation_status": "UNALLOCATED",
  "match_score": 95,
  "reason": "Vikram is a Senior Frontend Engineer with 8 years of experience,
             matching the seniority requirement. He is UNALLOCATED and his
             last project ended 120 days ago, fulfilling the temporal constraint.
             His role as Tech Lead on 'Storefront Platform Redesign' shows
             strong frontend leadership.",
  "matched_skill_names": ["React", "Next.js", "TypeScript", "Redux"],
  "evidence_snippets": [
    "Tech Lead — Storefront Platform Redesign",
    "8 years building production React applications"
  ]
}
""",
    )

    add_para(doc, "Call 4: JD Distillation (jd_distill_v1.py)", bold=True)
    add_para(doc, "Model: settings.GEMINI_MODEL_LIGHT (default gemini-2.5-flash). Temperature: 0.1, plain-text response (not JSON).")
    add_para(doc, "Purpose: take a verbose job description and compress it into ONE concise hiring query (≤200 chars). That query feeds the same /search ranker.")
    add_para(doc, "System prompt (verbatim):", italic=True)
    add_code_block(
        doc,
        """\
You are an expert technical recruiter. You take a verbose job description and
distill it into ONE concise hiring query (max 200 characters) that captures
the must-have skills, seniority, location, and any constraints.

Output ONLY the query sentence. No prose, no "Query:" prefix, no markdown.

Rules:
- Lead with seniority + role (e.g. "Senior backend engineer").
- List 2-4 must-have skills/technologies.
- Include hard constraints: city, years, domain expertise, availability.
- Drop fluff: "rockstar", "ninja", "fast-paced environment", "competitive salary".
- Drop nice-to-haves unless there are no hard requirements.
- Use the same phrasing a hiring manager would say out loud.
""",
    )
    add_para(doc, "Why this design matters: separating distillation from ranking means we can show HR what the AI heard before it ranks anything — a transparency win. It also keeps the ranker prompt unchanged.")

    add_heading(doc, "7.3 Prompt Engineering Principles We Followed", level=2)
    add_bullet(doc, "Structured JSON output, always. Every prompt declares a schema and forbids prose; safe_json_loads tolerates stray code fences but the prompt does the heavy lifting.")
    add_bullet(doc, "Few-shot examples for the hard prompts. Both extraction and inference ship with one full worked example. Search and JD-distillation use multiple short examples instead — the diversity of queries matters more than depth.")
    add_bullet(doc, "Evidence-by-default. Extraction requires an evidence quote per skill; search requires evidence_snippets per result. This is the anti-hallucination guardrail visible in the UI.")
    add_bullet(doc, "Explicit rules for edge cases. Verbatim proficiency mapping (architected→EXPERT). Hard date thresholds (90 days = quarter). Maximum counts (≤5 inferences, confidence ≥0.7). All ambiguity that matters is resolved in the prompt, not in code.")
    add_bullet(doc, "Two-tier model selection. Showcase model (2.5-pro) for the reasoning-heavy calls, light model (2.5-flash) for deterministic rewriting. Configurable in .env — no code change to swap.")
    add_bullet(doc, "Temporal context is injected, never hallucinated. The search prompt computes today/last_month/last_quarter on the server and ships them in the prompt, plus a precomputed days_since_last_project_end on every candidate.")
    add_bullet(doc, "Prompts live in code, versioned. extraction_v1.py, inference_v1.py, search_v1.py, jd_distill_v1.py — each file's name encodes its version, so a v2 prompt is a new file and a git diff away from an A/B test.")

    add_heading(doc, "7.4 Defensive Engineering Around AI", level=2)
    add_para(doc, "safe_json_loads (app/common/utils.py). Strips ```json fences and stray \"json\" labels before parsing. Defensive because Gemini occasionally ignores the \"no markdown\" instruction. If parsing still fails, the caller raises UpstreamAIError with a clear message.")
    add_para(doc, "_humanize_gemini_error (modules/ai/service.py). Maps raw google-genai exceptions to user-facing messages: 429 → \"Daily AI quota reached for the configured Gemini model. Try again after the quota resets, or switch GEMINI_MODEL_SHOWCASE in your .env to a different model.\" 503 → \"AI service is temporarily overloaded.\" 401/403 → \"Check that GOOGLE_API_KEY is valid.\" Every UpstreamAIError carries this message through, and search/service.py forwards str(exc) into the SSE event: error payload. The frontend reads parsed.message.")
    add_para(doc, "SSE streaming, the \"perceived\" way. The Gemini SDK is synchronous; we run generate_content_stream in a worker thread and forward chunks through an asyncio.Queue. _iter_complete_objects walks brace depth + string state to extract complete JSON elements as they arrive — even when an evidence_snippets value contains a brace. This gives the demo \"streaming reasoning\" with no extra dependency.")
    add_para(doc, "Inference degrades non-fatally. If the inference call errors or returns bad JSON, the resume still saves with explicit skills only. No user-facing failure for a non-critical AI call.")
    add_para(doc, "Best-effort Supabase Storage. If SUPABASE_URL is unset, _upload_to_supabase returns None and the extraction still runs. The \"View resume PDF\" link is the only UX casualty.")
    add_para(doc, "Production roadmap acknowledged. The README and this document both call out what would change at scale: embedding prefilter for >1k profiles, real Anthropic-style server-streaming if we move to Anthropic, audit-log dashboards on search_query_logs, prompt-quality golden tests.")

    add_heading(doc, "7.5 Why Our AI Is Not Just a Keyword Matcher", level=2)
    add_para(doc, "Concrete example. HR types \"find me someone with payment gateway experience\". A keyword matcher returns only profiles whose text literally contains \"payment gateway\".")
    add_para(doc, "SkillsHub's search prompt explicitly tells the model: \"payment gateway\" matches Razorpay, Stripe, PayU, Paytm, Square, Braintree. So a candidate whose resume says \"integrated Razorpay checkout serving 2M+ monthly transactions\" surfaces at the top, with matched_skill_names: [\"Razorpay\"] and an evidence snippet quoting that exact sentence.")
    add_para(doc, "Three more cases the prompt handles by design:")
    add_bullet(doc, "\"Real-time features\" → Socket.IO, WebSocket, WebRTC, server-sent events, Firebase RTDB. A candidate with \"built a live-collaboration whiteboard with WebSocket\" is a strong match.")
    add_bullet(doc, "\"Senior frontend folks who haven't been on a new project in the last quarter\" → seniority + frontend + days_since_last_project_end > 90 OR UNALLOCATED. The model returns Sneha (1199 days), Vikram (120 days), Neha (99 days) and cites the day counts in its reason.")
    add_bullet(doc, "\"ML\" → PyTorch, TensorFlow, scikit-learn, Hugging Face, LangChain. The prompt makes the equivalence explicit so we don't depend on the model spontaneously knowing.")
    add_para(doc, "Result: every result card carries reasoning (\"He architected the Razorpay integration serving 2M+ transactions/month\"), not just a name. Judges can read the reasoning and verify the match against the evidence quote without trusting the model blindly.")

    add_page_break(doc)

    # ─────────────────────────────────────────────────────────────
    # 8. SECURITY & ARCHITECTURE CHOICES
    # ─────────────────────────────────────────────────────────────
    add_heading(doc, "8. Security & Architecture Choices", level=1)
    add_bullet(doc, "JWT auth (HS256), 24-hour expiry by default (JWT_EXPIRATION_MINUTES=1440). Custom claims: sub, role, iat, exp. Generated via create_access_token, validated via decode_access_token (jose).")
    add_bullet(doc, "Passwords hashed with bcrypt via passlib (CryptContext schemes=[\"bcrypt\"]). We never log raw passwords; the only place plaintext exists is the request body, which is excluded from access logs by FastAPI's default.")
    add_bullet(doc, "Role-based access at every protected endpoint. Two dependency-injectable guards: get_current_user (verifies the JWT and loads the User row) and require_role(*roles). The shorthand require_admin = require_role(UserRole.ADMIN) is applied at the router level for the entire /review module.")
    add_bullet(doc, "Ownership checks at the service layer. EmployeeService.ensure_can_view and ensure_can_edit gate access — USERs can only see/edit their own employee row; ADMINs can do both.")
    add_bullet(doc, "CORS configured from CORS_ORIGINS env var; comma-separated origins parsed at startup. Defaults to localhost:3000 only.")
    add_bullet(doc, "Auth rate limit: 10 req/min/IP (configurable via AUTH_RATE_LIMIT_PER_MINUTE). In-memory sliding-window — fine for hackathon scope, would be Redis-backed in production.")
    add_bullet(doc, "Failed login returns 401 with a generic \"Invalid email or password.\" message — never reveals whether the email exists.")
    add_bullet(doc, "HR signup is invite-coded. payload.invite_code is compared to settings.HR_INVITE_CODE with hmac.compare_digest (constant-time, defeats timing side channels). Wrong code → 403 with generic message.")
    add_bullet(doc, "Frontend middleware is a UX guard, not a security boundary. It decodes the JWT to route by role early — but the backend re-verifies the signature on every protected request.")
    add_bullet(doc, "Environment variables validated at startup via pydantic-settings. DATABASE_URL, JWT_SECRET, GOOGLE_API_KEY are required — the app fails to boot without them. .env is gitignored; .env.example commits non-secret keys with placeholder values.")
    add_bullet(doc, "Standard error envelope. Every error response is shaped { success, statusCode, message, timestamp, path } — see app/common/filters.py. Frontend axios interceptor reads .message.")
    add_bullet(doc, "Request IDs on every response. RequestIdMiddleware generates a uuid4 (or honours x-request-id), logs the request with elapsed_ms, and writes the id back as a response header.")
    add_para(doc, "What we'd add for production:", bold=True)
    add_bullet(doc, "httpOnly cookie auth instead of localStorage. localStorage is fine for the demo; XSS exposure is the production concern.")
    add_bullet(doc, "Distributed rate limiting (slowapi + Redis) so the limit holds across replicas.")
    add_bullet(doc, "SSO integration (Google / Okta / Azure AD) for enterprise.")
    add_bullet(doc, "Audit logging on every AI call (which prompt version, which model, which user, full input/output) so prompt regressions are diagnosable.")
    add_bullet(doc, "Prompt-injection defences on resume text (it's user-controlled). Phase 1 trusts the resume parser to ignore embedded \"ignore previous instructions\" lines — production would strip those defensively.")
    add_bullet(doc, "Resume re-extraction job queue — re-run when the prompt version bumps.")

    # ─────────────────────────────────────────────────────────────
    # 9. PHASE 1 vs FUTURE PHASES
    # ─────────────────────────────────────────────────────────────
    add_heading(doc, "9. What's in Phase 1 (vs Future Phases)", level=1)
    add_table(
        doc,
        ["In Phase 1", "Future Phases"],
        [
            ["JWT auth with HR + Employee roles", "SSO / Active Directory / Google login"],
            ["Email + password signup with invite-coded HR", "Magic-link / OTP signup"],
            ["AI resume extraction (gemini-2.5-pro)", "OCR for scanned (image-only) PDFs"],
            ["Related-skill inference with confidence + reason", "Multi-source skill enrichment (GitHub, public talks)"],
            ["NL search with streamed ranking + reasoning", "Embedding prefilter for >1 000 profiles"],
            ["JD-to-shortlist (paste a JD, get a ranked list)", "Live ATS sync — JDs come from Greenhouse/Workday"],
            ["Team Builder (single-call team assembly)", "Capacity-aware team builder with calendar/allocation feed"],
            ["Single-reviewer queue (HR approves)", "Multi-reviewer workflows, delegation, SLAs"],
            ["Audit log of every search (search_query_logs)", "Prompt-quality dashboard built on the audit log"],
            ["Versioned prompts in code (v1)", "Prompt A/B testing + golden-set regression suite"],
            ["Best-effort Supabase Storage for PDFs", "Direct-to-S3 upload with presigned URLs"],
            ["In-memory auth rate limit", "Redis-backed distributed rate limiting"],
        ],
        col_widths=[3.1, 3.1],
    )

    add_page_break(doc)

    # ─────────────────────────────────────────────────────────────
    # 10. DEMONSTRATING SKILLSHUB
    # ─────────────────────────────────────────────────────────────
    add_heading(doc, "10. Demonstrating SkillsHub to Judges", level=1)
    add_para(doc, "Written for the team member running the live demo. Plain language. Step-by-step. Assume you might forget what to click.")

    add_heading(doc, "10.1 Pre-Flight Checklist (10 minutes before judging)", level=2)
    add_bullet(doc, "Backend up: from /backend, run `uvicorn app.main:app --reload --port 8000`. Open http://localhost:8000/docs — Swagger should render.")
    add_bullet(doc, "Frontend up: from /frontend, run `npm run dev`. Open http://localhost:3000/login.")
    add_bullet(doc, "Reset the demo DB: from /backend, run `python -m scripts.seed_database`. You should see \"Upserted HR user → hr@skillshub.demo\" and 15 profile lines.")
    add_bullet(doc, "Have two browser tabs ready: one logged in as HR (/search), one in incognito for the employee flow.")
    add_bullet(doc, "Have a real PDF resume on your desktop. Use one with clear text — not a scanned image.")
    add_bullet(doc, "Check Gemini quota: from /backend, run `python -m scripts._stream_probe`. If it returns 429, switch GEMINI_MODEL_SHOWCASE in .env to gemini-2.5-pro (or a different model with fresh quota) and restart uvicorn.")
    add_bullet(doc, "Network: open Network tab in browser devtools — useful for showing the SSE stream live if you have time.")

    add_heading(doc, "10.2 The 3-Minute Demo Script", level=2)

    add_para(doc, "[0:00 – 0:15] Open with the hook", bold=True)
    add_para(doc, "WHAT TO SAY: \"In a 200-person dev team, when HR needs someone who knows React and has built payment integrations, the answer lives in spreadsheets and in people's heads. We built SkillsHub to fix that.\"")
    add_para(doc, "WHAT TO DO: Open the HR tab on /search.")

    add_para(doc, "[0:15 – 0:50] Show search magic", bold=True)
    add_para(doc, "WHAT TO SAY: \"Watch a real query. I'll ask in plain English.\"")
    add_para(doc, "WHAT TO DO: Type `Find me a backend dev in Pune with 3 years of Java and payment gateway experience` and press Enter.")
    add_para(doc, "WHAT JUDGES SEE: Cards stream in one by one. Score badges fill. Matched skills are highlighted.")
    add_para(doc, "WHAT TO POINT AT: \"See the top card cites the specific Razorpay project? That's not keyword match — the prompt explicitly maps 'payment gateway' to Razorpay, Stripe, PayU. It's reasoning over the resume.\"")

    add_para(doc, "[0:50 – 1:20] Show temporal reasoning", bold=True)
    add_para(doc, "WHAT TO SAY: \"And it understands time, not just skills.\"")
    add_para(doc, "WHAT TO DO: Clear, then type `Senior frontend folks who haven't been on a new project in the last quarter.`")
    add_para(doc, "WHAT JUDGES SEE: Sneha, Vikram, Neha bubble up — UNALLOCATED, 1199/120/99 days since last project.")
    add_para(doc, "WHAT TO POINT AT: \"Notice the reason cites the day count — 120 days ago, 99 days ago. The prompt gets today's date as context plus a precomputed days_since_last_project_end on every candidate. No date hallucination.\"")

    add_para(doc, "[1:20 – 1:45] Show JD-to-Shortlist", bold=True)
    add_para(doc, "WHAT TO SAY: \"Same idea but starting from a real job description.\"")
    add_para(doc, "WHAT TO DO: Click the From JD tab. Paste a 5-line JD. Click Find Shortlist.")
    add_para(doc, "WHAT JUDGES SEE: A distilled query pill appears first (\"Senior backend engineer in Pune with 5+ years Java/Spring Boot and payment gateway…\") — that's a fast Gemini call. Then the same ranked cards stream in.")
    add_para(doc, "WHAT TO POINT AT: \"Two-stage SSE — first event is the distilled query so you see exactly what the AI heard. Then the same ranker runs.\"")

    add_para(doc, "[1:45 – 2:30] Switch to Employee — upload flow", bold=True)
    add_para(doc, "WHAT TO SAY: \"This is how a profile gets in. Employee account, drag a resume.\"")
    add_para(doc, "WHAT TO DO: Switch to the incognito tab. Login as Employee (one click). You land on /upload. Drop the PDF.")
    add_para(doc, "WHAT JUDGES SEE: 4-step animated progress (Upload → Read PDF → Extract with Gemini → Infer related skills). After ~6 seconds the profile editor fades in with skills + an amber Inferred Skills section.")
    add_para(doc, "WHAT TO POINT AT: \"Hover over an inferred skill — the tooltip shows the confidence and the reason. Next.js → React at 98% because Next.js is built on React.\" Click Send for Review.")

    add_para(doc, "[2:30 – 2:55] Back to HR — approve + verify", bold=True)
    add_para(doc, "WHAT TO SAY: \"HR sees the new profile, approves it, and it's live.\"")
    add_para(doc, "WHAT TO DO: Switch to the HR tab. Click Review in the nav. Click the new item. Hit Approve.")
    add_para(doc, "WHAT TO POINT AT: \"Now go back to /search and search for the new person's signature skill — they appear. Human-in-the-loop, no AI output ships to a hiring manager unreviewed.\"")

    add_para(doc, "[2:55 – 3:00] Close", bold=True)
    add_para(doc, "WHAT TO SAY: \"Two screens, one AI spine, full audit log. Questions?\"")

    add_heading(doc, "10.3 Likely Judge Questions & Crisp Answers", level=2)

    add_para(doc, "Q: How does the proficiency inference work?", bold=True)
    add_para(doc, "A: Verb-based rules in the extraction prompt. \"architected/led/owned\" → EXPERT, \"built/developed/implemented\" → INTERMEDIATE, \"used/familiar with/exposed to\" → NOVICE. The full ruleset is in backend/app/modules/ai/prompts/extraction_v1.py.")

    add_para(doc, "Q: How would this scale to 10 000 profiles?", bold=True)
    add_para(doc, "A: Add an embedding prefilter — Supabase pgvector or a managed embedding API. Retrieve top-50 by cosine similarity, then run the same single-call ranking prompt on those 50. The prompt is unchanged; only the candidate set shrinks.")

    add_para(doc, "Q: What about prompt injection in resumes?", bold=True)
    add_para(doc, "A: Real concern. Phase 1 trusts the parser model to ignore \"ignore previous instructions\" lines (Gemini's instruction-following is robust enough for hackathon scope). Production would strip suspicious patterns from raw_text before the call, run extraction inside a constrained system role, and validate that the output schema matches.")

    add_para(doc, "Q: Why Gemini and not Claude?", bold=True)
    add_para(doc, "A: Free-tier daily quota during development, 2.5-pro's reasoning quality on structured JSON is competitive, and the google-genai SDK's streaming generator works cleanly with our SSE pipeline. The AI client is one file — swapping to Anthropic SDK is roughly 30 lines.")

    add_para(doc, "Q: What about scanned (image-only) PDFs?", bold=True)
    add_para(doc, "A: Phase 1 rejects them with a clear error. _pdf_to_text returns the extracted text length; if <80 chars it raises \"PDF text was too short to parse — is this a scanned image?\". Phase 2 would wire in an OCR pass (Tesseract or Google Vision).")

    add_para(doc, "Q: How do you know it's not hallucinating?", bold=True)
    add_para(doc, "A: Every result card carries evidence_snippets — exact quotes from the candidate's resume. Every extracted skill carries an evidence quote. We can audit any decision back to the source text.")

    add_para(doc, "Q: What happens when the AI service is down?", bold=True)
    add_para(doc, "A: The user sees a specific, humanised error: \"Daily AI quota reached…\" or \"AI service is temporarily overloaded…\". Inference errors degrade non-fatally — explicit skills still save. The error pipeline is _humanize_gemini_error in modules/ai/service.py.")

    add_para(doc, "Q: Why no embedding search at all?", bold=True)
    add_para(doc, "A: At 15–30 profiles the LLM call dominates cost, and an embedding prefilter introduces a recall risk (right candidate gets cut). With more profiles we'd add it. The point is the single-call ranking is the higher-quality baseline.")

    add_heading(doc, "10.4 If Something Breaks Mid-Demo", level=2)
    add_bullet(doc, "Search returns the humanised \"Daily AI quota reached\" message → say \"Free-tier quota — let me show you the same query against our pre-recorded run.\" Have a screenshot of a successful search ready.")
    add_bullet(doc, "Resume upload spins forever → hit Cancel, say \"Cold start on the free-tier API — let me show you a profile that's already extracted\" and click into a directory profile instead.")
    add_bullet(doc, "Backend 500 → check the uvicorn console for the traceback. If it's a Supabase Storage timeout, the resume still saved — refresh the upload page and the profile editor will appear.")
    add_bullet(doc, "Search corpus empty → run `python -m scripts.seed_database` to reseed; ten seconds.")
    add_bullet(doc, "Frontend cache weird (loader stuck, layout missing) → Ctrl+Shift+R for hard reload. If still broken, `rm -rf .next && npm run dev`.")
    add_bullet(doc, "Login fails → check that you copied the demo password correctly; it's `demo123`. If still failing, look at the Network tab — 429 means you exceeded the auth rate limit (10/min), wait a minute.")

    add_page_break(doc)

    # ─────────────────────────────────────────────────────────────
    # APPENDIX A — Endpoints
    # ─────────────────────────────────────────────────────────────
    add_heading(doc, "Appendix A: API Endpoint Reference", level=1)
    add_para(doc, "Full Swagger at /docs. All non-auth, non-meta endpoints require Bearer JWT.")
    add_table(
        doc,
        ["Method", "Path", "Auth", "Purpose"],
        [
            ["GET", "/", "Public", "Health check."],
            ["POST", "/auth/register/employee", "Public", "Open employee signup. Creates User + stub Employee."],
            ["POST", "/auth/register/hr", "Public", "Invite-coded HR signup."],
            ["POST", "/auth/login", "Public", "Returns access_token + user + employee_id."],
            ["GET", "/auth/me", "Authenticated", "Current user; used to rehydrate on refresh."],
            ["GET", "/users/me", "Authenticated", "Same as /auth/me but user-only payload."],
            ["GET", "/employees", "Authenticated", "Directory list with filters + pagination."],
            ["GET", "/employees/me", "Authenticated", "Logged-in employee's own profile."],
            ["GET", "/employees/{id}", "Authenticated", "Single profile (ownership-checked for USER)."],
            ["PATCH", "/employees/{id}", "Authenticated", "Update an employee (ownership rules apply)."],
            ["PUT", "/employees/{id}/skills", "Authenticated", "Replace the skill set."],
            ["PUT", "/employees/{id}/projects", "Authenticated", "Replace the project set."],
            ["POST", "/resumes/upload", "Authenticated", "PDF upload + AI extraction + inference pipeline."],
            ["GET", "/review/queue", "ADMIN", "Pending review items."],
            ["GET", "/review/{id}", "ADMIN", "Review item + the employee profile."],
            ["POST", "/review/{id}/approve", "ADMIN", "Approve a pending item."],
            ["POST", "/review/{id}/reject", "ADMIN", "Reject a pending item."],
            ["POST", "/review/{id}/edit-and-approve", "ADMIN", "Mark as edited-and-approved (after HR saved edits)."],
            ["POST", "/search", "ADMIN", "Natural-language search (SSE)."],
            ["POST", "/search/jd", "ADMIN", "Paste-a-JD search (two-stage SSE: distill then rank)."],
            ["POST", "/search/sync", "ADMIN", "Non-streaming search variant for debugging."],
            ["POST", "/search/team", "ADMIN", "Team builder (single-call team assembly)."],
        ],
        col_widths=[0.8, 2.4, 1.1, 2.4],
    )

    # ─────────────────────────────────────────────────────────────
    # APPENDIX B — Env vars
    # ─────────────────────────────────────────────────────────────
    add_heading(doc, "Appendix B: Environment Variables", level=1)
    add_table(
        doc,
        ["Variable", "Purpose", "Example (secrets masked)"],
        [
            ["APP_PORT", "Bind port for uvicorn.", "8000"],
            ["APP_ENV", "development | production.", "development"],
            ["APP_NAME", "Used in OpenAPI title and logs.", "SkillsHub"],
            ["LOG_LEVEL", "DEBUG/INFO/WARNING/ERROR.", "INFO"],
            ["DATABASE_URL", "Async Supabase Postgres URL.", "postgresql+asyncpg://postgres.YOURREF:xxxxxxxx@aws-0-ap-south-1.pooler.supabase.com:5432/postgres"],
            ["DATABASE_URL_SYNC", "Sync URL for Alembic.", "postgresql+psycopg2://postgres.YOURREF:xxxxxxxx@aws-0-ap-south-1.pooler.supabase.com:5432/postgres"],
            ["JWT_SECRET", "Signing key (HS256). Required.", "xxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxx"],
            ["JWT_ALGORITHM", "Signing algorithm.", "HS256"],
            ["JWT_EXPIRATION_MINUTES", "Token lifetime.", "1440"],
            ["SUPABASE_URL", "Supabase project URL (optional).", "https://YOURREF.supabase.co"],
            ["SUPABASE_ANON_KEY", "Optional. For client-side reads.", "eyJxxx..."],
            ["SUPABASE_SERVICE_ROLE_KEY", "Required only if uploading resumes to Storage.", "eyJxxx..."],
            ["SUPABASE_STORAGE_BUCKET", "Bucket name for resumes.", "resumes"],
            ["GOOGLE_API_KEY", "Gemini API key. Required.", "AIzaSyxxx..."],
            ["GEMINI_MODEL_SHOWCASE", "Reasoning-heavy model (extraction, search).", "gemini-2.5-pro"],
            ["GEMINI_MODEL_LIGHT", "Fast model (inference, JD distill).", "gemini-2.5-flash"],
            ["CORS_ORIGINS", "Comma-separated allowed origins.", "http://localhost:3000,http://127.0.0.1:3000"],
            ["SEED_HR_EMAIL", "Demo HR account.", "hr@skillshub.demo"],
            ["SEED_HR_PASSWORD", "Demo HR password.", "demo123"],
            ["SEED_EMP_EMAIL", "Demo employee account.", "ravi@skillshub.demo"],
            ["SEED_EMP_PASSWORD", "Demo employee password.", "demo123"],
            ["HR_INVITE_CODE", "Required for HR signup.", "SKILLSHUB-HR-2026"],
            ["ALLOW_EMPLOYEE_SIGNUP", "Master switch.", "true"],
            ["ALLOW_HR_SIGNUP", "Master switch.", "true"],
            ["PASSWORD_MIN_LENGTH", "Minimum length on signup.", "8"],
            ["AUTH_RATE_LIMIT_PER_MINUTE", "Per-IP /auth/* limit (0 disables).", "10"],
        ],
        col_widths=[2.0, 2.4, 2.2],
    )

    # ─────────────────────────────────────────────────────────────
    # APPENDIX C — Team setup
    # ─────────────────────────────────────────────────────────────
    add_heading(doc, "Appendix C: Team Setup — Zero to Running in 10 Minutes", level=1)
    add_para(doc, "1) Clone and enter the repo.", bold=True)
    add_code_block(
        doc,
        """\
git clone <repo-url> skillshub-ai
cd skillshub-ai
""",
    )
    add_para(doc, "2) Set up the backend environment.", bold=True)
    add_code_block(
        doc,
        """\
cd backend
python -m venv .venv
.venv\\Scripts\\activate         # Windows
# source .venv/bin/activate     # macOS/Linux

pip install -r requirements.txt
cp .env.example .env
# Fill DATABASE_URL, DATABASE_URL_SYNC, JWT_SECRET, GOOGLE_API_KEY.
# Optional: SUPABASE_* for resume storage.
""",
    )
    add_para(doc, "3) Migrate + seed the database.", bold=True)
    add_code_block(
        doc,
        """\
alembic upgrade head
python -m scripts.seed_database
""",
    )
    add_para(doc, "You should see \"Upserted HR user → hr@skillshub.demo\" and 15 profile lines.")
    add_para(doc, "4) Start the backend.", bold=True)
    add_code_block(
        doc,
        """\
uvicorn app.main:app --reload --port 8000
# Open http://localhost:8000/docs — Swagger UI should render.
""",
    )
    add_para(doc, "5) Set up + start the frontend.", bold=True)
    add_code_block(
        doc,
        """\
cd ../frontend
npm install
# Create .env.local with:
#   NEXT_PUBLIC_API_BASE_URL=http://localhost:8000
npm run dev
# Open http://localhost:3000.
""",
    )
    add_para(doc, "6) Log in with the seeded accounts.", bold=True)
    add_table(
        doc,
        ["Role", "Email", "Password"],
        [
            ["HR / ADMIN", "hr@skillshub.demo", "demo123"],
            ["Employee", "ravi@skillshub.demo", "demo123"],
        ],
        col_widths=[1.4, 2.5, 1.4],
    )
    add_para(doc, "Both accounts are available as one-click buttons on /login. Re-running the seed script restores them at any time.")

    # ── Save ──────────────────────────────────────────────────────
    doc.save(str(OUT_PATH))
    print(f"WROTE: {OUT_PATH}")


if __name__ == "__main__":
    build()
